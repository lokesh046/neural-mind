import os
import base64
import io

def generate_assets():
    print("Generating RAG binary assets...")
    text = (
        "Retrieval-Augmented Generation (RAG) is a technique that bridges the gap between static LLM "
        "knowledge and dynamic private databases. First, the Ingestion layer reads PDF or Word files. "
        "Next, the Parser converts these binary streams to plain text. Then, the Chunking step splits "
        "text using character boundaries and sliding overlaps. The Embedding model projects chunks "
        "into high-dimensional vectors. FAISS stores and indexes these vectors. Finally, the Retriever "
        "queries FAISS to retrieve context and stuff it into the LLM prompt."
    )
    
    # 1. Generate PDF using reportlab
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    
    pdf_buffer = io.BytesIO()
    doc = SimpleDocTemplate(pdf_buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    story = [
        Paragraph("<b>RAG Technical Specification</b>", styles['Title']),
        Spacer(1, 12),
        Paragraph(text, styles['Normal'])
    ]
    doc.build(story)
    pdf_bytes = pdf_buffer.getvalue()
    pdf_b64 = base64.b64encode(pdf_bytes).decode('utf-8')
    print("Generated PDF Base64 string.")
    
    # 2. Generate DOCX using python-docx
    import docx
    doc_docx = docx.Document()
    doc_docx.add_heading('RAG Technical Specification', 0)
    doc_docx.add_paragraph(text)
    docx_buffer = io.BytesIO()
    doc_docx.save(docx_buffer)
    docx_bytes = docx_buffer.getvalue()
    docx_b64 = base64.b64encode(docx_bytes).decode('utf-8')
    print("Generated DOCX Base64 string.")
    
    # 3. Cache the model
    print("Pre-downloading and caching SentenceTransformer('all-MiniLM-L6-v2')...")
    from sentence_transformers import SentenceTransformer
    model = SentenceTransformer('all-MiniLM-L6-v2')
    print("Model cached successfully!")
    
    # Write base64 strings to a python file that we can easily import in seed.py
    output_path = r"c:\Users\lokes\OneDrive\Desktop\tensortonic_clone\backend\rag_assets_data.py"
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(f'PDF_BASE64 = """{pdf_b64}"""\n\n')
        f.write(f'DOCX_BASE64 = """{docx_b64}"""\n\n')
        f.write(f'SAMPLE_TEXT = """{text}"""\n')
    print(f"Saved base64 data strings to {output_path}")

if __name__ == "__main__":
    generate_assets()
